#!/usr/bin/env python3
"""
Fill Word document templates with data and convert them to PDF.
"""

from __future__ import annotations

import os
import shutil
import subprocess
import tempfile
from zipfile import ZIP_DEFLATED, ZipFile
from typing import Dict, Iterable
from xml.sax.saxutils import escape

from docx import Document

from create_word_templates import create_pressure_vessel_template, create_safety_valve_template


SAFETY_VALVE_TEMPLATE = "templates/Safety_Valve_Client_Template.docx"
PRESSURE_VESSEL_TEMPLATE = "templates/Pressure_Vessel_Client_Template.docx"


def ensure_safety_valve_template(refresh: bool = False) -> None:
    """Create the Safety Valve template when missing or refresh is requested."""
    if refresh or not os.path.exists(SAFETY_VALVE_TEMPLATE):
        create_safety_valve_template()


def ensure_pressure_vessel_template(refresh: bool = False) -> None:
    """Create the Pressure Vessel template when missing or refresh is requested."""
    if refresh or not os.path.exists(PRESSURE_VESSEL_TEMPLATE):
        create_pressure_vessel_template()


def _iter_all_paragraphs(doc: Document) -> Iterable:
    seen = set()

    for paragraph in doc.paragraphs:
        paragraph_id = id(paragraph)
        if paragraph_id not in seen:
            seen.add(paragraph_id)
            yield paragraph

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph_id = id(paragraph)
                    if paragraph_id not in seen:
                        seen.add(paragraph_id)
                        yield paragraph


def _replace_paragraph_text(paragraph, replacements: Dict[str, str]) -> bool:
    text = paragraph.text
    updated = text

    for placeholder, value in replacements.items():
        if placeholder in updated:
            updated = updated.replace(placeholder, value)

    if updated == text:
        return False

    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(updated)

    return True


def fill_template(template_path: str, output_path: str, data: Dict[str, str]) -> str:
    """Fill a Word template with data."""
    replacements = {f"{{{{{key}}}}}": escape(str(value)) for key, value in data.items()}
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_output = os.path.join(temp_dir, "filled.docx")
        with ZipFile(template_path, "r") as source_zip, ZipFile(temp_output, "w", ZIP_DEFLATED) as target_zip:
            for item in source_zip.infolist():
                content = source_zip.read(item.filename)
                if item.filename.endswith(".xml"):
                    text = content.decode("utf-8")
                    for placeholder, value in replacements.items():
                        text = text.replace(placeholder, value)
                    content = text.encode("utf-8")
                target_zip.writestr(item, content)
        shutil.copyfile(temp_output, output_path)
    return output_path


def _convert_with_soffice(docx_path: str, pdf_path: str) -> bool:
    output_dir = os.path.dirname(pdf_path) or "."
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path],
        check=True,
        capture_output=True,
    )

    generated_pdf = os.path.join(output_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
    if os.path.abspath(generated_pdf) != os.path.abspath(pdf_path) and os.path.exists(generated_pdf):
        shutil.move(generated_pdf, pdf_path)

    return os.path.exists(pdf_path)


def _convert_with_word_com(docx_path: str, pdf_path: str) -> bool:
    import comtypes.client

    word = comtypes.client.CreateObject("Word.Application")
    word.Visible = False
    document = None

    try:
        document = word.Documents.Open(os.path.abspath(docx_path))
        document.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
    finally:
        if document is not None:
            document.Close(False)
        word.Quit()

    return os.path.exists(pdf_path)


def convert_to_pdf(docx_path: str, pdf_path: str) -> bool:
    """Convert Word document to PDF using LibreOffice or Microsoft Word."""
    try:
        return _convert_with_soffice(docx_path, pdf_path)
    except (subprocess.CalledProcessError, FileNotFoundError):
        pass

    try:
        return _convert_with_word_com(docx_path, pdf_path)
    except Exception:
        return False


def _generate_from_template(template_path: str, data: Dict[str, str], output_pdf: str) -> str:
    filled_docx = output_pdf[:-4] + ".docx" if output_pdf.lower().endswith(".pdf") else output_pdf + ".docx"

    if os.path.exists(filled_docx):
        os.remove(filled_docx)
    if os.path.exists(output_pdf):
        os.remove(output_pdf)

    fill_template(template_path, filled_docx, data)

    if convert_to_pdf(filled_docx, output_pdf):
        if os.path.exists(filled_docx):
            os.remove(filled_docx)
        return output_pdf

    return filled_docx


def map_safety_valve_form_data(data: Dict[str, str]) -> Dict[str, str]:
    """Map desktop app field names to template placeholders."""
    return {
        "CLIENT_COMPANY": data.get("clientCompany", ""),
        "CLIENT_ADDRESS": data.get("clientAddress", ""),
        "REPORT_DATE": data.get("reportDate", ""),
        "JOB_NUMBER": data.get("jobNumber", ""),
        "REFERENCE": data.get("reference", ""),
        "TEST_NUMBER": data.get("testNumber", ""),
        "CONTACT": data.get("contact", ""),
        "PO_REFERENCE": data.get("poReference", ""),
        "TEST_DATE": data.get("testDate", ""),
        "VALVE_SIZE": data.get("valveSize", ""),
        "MANUFACTURER": data.get("manufacturer", ""),
        "MODEL": data.get("model", ""),
        "DISCHARGE_TYPE": data.get("dischargeType", "Atmosphere"),
        "SET_PRESSURE": data.get("setPressure", ""),
        "VALVE_SERIAL_NO": data.get("valveSerialNo", ""),
        "BLOWDOWN_TYPE": data.get("blowdownType", "Non-Adjustable"),
        "TEST_GAUGE_ID": data.get("testGaugeId", "N/A"),
        "GAUGE_MODEL": data.get("gaugeModel", "Druck DPI104"),
        "GAUGE_SERIAL_NO": data.get("gaugeSerialNo", "6113016"),
        "CERTIFICATION_NO": data.get("certificationNo", "N/A"),
        "RECALIBRATION_DUE": data.get("recalibrationDue", ""),
        "TEST_MEDIUM": data.get("testMedium", "Nitrogen"),
        "AVG_LIFT_PRESSURE": data.get("avgLiftPressure", ""),
        "AVG_RESEAT_PRESSURE": data.get("avgReseatPressure", ""),
        "FLOW_RATE": data.get("flowRate", ""),
        "BLOWDOWN_PERCENT": data.get("blowdownPercent", ""),
        "ALLOWABLE_BLOWDOWN": data.get("allowableBlowdown", "<15"),
        "LEAK_TEST_RESULT": data.get("leakTest", "PASS"),
        "SET_PRESSURE_TEST_RESULT": data.get("setPressureTest", "PASS"),
        "BLOWDOWN_TEST_RESULT": data.get("blowdownTest", "PASS"),
        "OVERALL_RESULT": data.get("overallResult", "PASS"),
        "COMMENTS": data.get("comments", ""),
        "TECHNICIAN_NAME": data.get("technicianName", "Glenn Scatchard"),
        "SIGNATURE_DATE": data.get("signatureDate", ""),
    }


def map_pressure_vessel_form_data(data: Dict[str, str]) -> Dict[str, str]:
    """Map desktop app field names to template placeholders."""
    return {
        "CLIENT_COMPANY": data.get("clientCompany", ""),
        "CLIENT_ADDRESS": data.get("clientAddress", ""),
        "REPORT_DATE": data.get("reportDate", ""),
        "JOB_NUMBER": data.get("jobNumber", ""),
        "REFERENCE": data.get("reference", ""),
        "MANUFACTURER": data.get("manufacturer", ""),
        "MODEL": data.get("model", ""),
        "SERIAL_NO": data.get("serialNo", ""),
        "VESSEL_MANUFACTURER": data.get("vesselManufacturer", data.get("manufacturer", "")),
        "VESSEL_SERIAL_NO": data.get("vesselSerialNo", data.get("serialNo", "")),
        "TEST_NUMBER": data.get("testNumber", ""),
        "CONTACT": data.get("contact", ""),
        "PO_REFERENCE": data.get("poReference", ""),
        "ASSET_NO": data.get("assetNo", ""),
        "TEST_DATE": data.get("testDate", ""),
        "UNIT_TYPE": data.get("unitType", ""),
        "MANUFACTURED_DATE": data.get("manufacturedDate", ""),
        "LOCATION": data.get("location", "Itinerant"),
        "DESIGN_CODE": data.get("designCode", "AS 1210-3"),
        "CORROSION_ALLOWANCE": data.get("corrosionAllowance", ""),
        "VESSEL_TYPE": data.get("vesselType", ""),
        "DESIGN_REG_NO": data.get("designRegNo", ""),
        "WORKSAFE_REG_NO": data.get("worksafeRegNo", ""),
        "COMMISSIONING_DATE": data.get("commissioningDate", ""),
        "DESIGN_PRESSURE": data.get("designPressure", ""),
        "OPERATING_PRESSURE": data.get("operatingPressure", ""),
        "TEST_PRESSURE": data.get("testPressure", ""),
        "HAZARD_LEVEL": data.get("hazardLevel", ""),
        "TANK_CAPACITY": data.get("tankCapacity", ""),
        "HYDROSTATIC_TEST_DATE": data.get("hydrostaticTestDate", ""),
        "DESIGN_TEMP": data.get("designTemp", ""),
        "AMBIENT_TEMP": data.get("ambientTemp", ""),
        "SHELL_LENGTH": data.get("shellLength", ""),
        "SHELL_DIAMETER": data.get("shellDiameter", ""),
        "NEXT_INSPECTION_DUE": data.get("nextInspectionDue", ""),
        "OVERALL_RESULT": data.get("overallResult", "PASS"),
        "INSPECTION_TYPE": data.get("inspectionType", ""),
        "LHS_LID_SPEC": data.get("lhsLidSpec", ""),
        "LHS_LID_ACTUAL": data.get("lhsLidActual", ""),
        "RHS_BASE_SPEC": data.get("rhsBaseSpec", ""),
        "RHS_BASE_ACTUAL": data.get("rhsBaseActual", ""),
        "SHELL_SPEC": data.get("shellSpec", ""),
        "SHELL_ACTUAL": data.get("shellActual", ""),
        "EXTERNAL_SURFACE_CONDITION": data.get("externalSurfaceCondition", ""),
        "EXTERNAL_SURFACE_COMMENTS": data.get("externalSurfaceComments", ""),
        "INTERNAL_SURFACE_CONDITION": data.get("internalSurfaceCondition", ""),
        "INTERNAL_SURFACE_COMMENTS": data.get("internalSurfaceComments", ""),
        "CORROSION_OBSERVED": data.get("corrosionObserved", ""),
        "CORROSION_LOCATION": data.get("corrosionLocation", ""),
        "DENTS_DEFORMATION": data.get("dentsDeformation", ""),
        "DENTS_DETAILS": data.get("dentsDetails", ""),
        "PRESSURE_GAUGE_WORKING": data.get("pressureGaugeWorking", ""),
        "PAINT_COATING_CONDITION": data.get("paintCoatingCondition", ""),
        "BASE_MOUNTS_CONDITION": data.get("baseMountsCondition", ""),
        "BASE_MOUNTS_COMMENTS": data.get("baseMountsComments", ""),
        "NAMEPLATE_STATUS": data.get("nameplateStatus", ""),
        "DRAIN_VALVE_OPERATION": data.get("drainValveOperation", ""),
        "SAFETY_VALVE_FITTED": data.get("safetyValveFitted", ""),
        "SAFETY_VALVE_REPORT_NO": data.get("safetyValveReportNo", ""),
        "SAFETY_VALVE_DETAILS": data.get("safetyValveDetails", ""),
        "OVERALL_VESSEL_CONDITION": data.get("overallVesselCondition", ""),
        "DESIGN_REPORTS_SIGHTED": data.get("designReportsSighted", ""),
        "ASSOCIATED_DOCUMENTATION": data.get("associatedDocumentation", ""),
        "COMMENTS": data.get("comments", ""),
        "TECHNICIAN_NAME": data.get("inspectorName", "Glenn Scatchard"),
        "SIGNATURE_DATE": data.get("signatureDate", ""),
    }


def generate_safety_valve_pdf(data: Dict[str, str], output_pdf: str) -> str:
    """Generate a Safety Valve PDF from the client-based Word template."""
    ensure_safety_valve_template(refresh=True)
    return _generate_from_template(SAFETY_VALVE_TEMPLATE, data, output_pdf)


def generate_pressure_vessel_pdf(data: Dict[str, str], output_pdf: str) -> str:
    """Generate a Pressure Vessel PDF from the client-based Word template."""
    ensure_pressure_vessel_template(refresh=True)
    return _generate_from_template(PRESSURE_VESSEL_TEMPLATE, data, output_pdf)


SAFETY_VALVE_SAMPLE_DATA = {
    "CLIENT_COMPANY": "ABC Manufacturing Pty Ltd",
    "CLIENT_ADDRESS": "123 Industrial Road, Perth WA 6000",
    "REPORT_DATE": "2026-03-30",
    "JOB_NUMBER": "WAC-2024-001",
    "REFERENCE": "VALVE-INSPECT-001",
    "TEST_NUMBER": "TV-001",
    "CONTACT": "John Smith",
    "PO_REFERENCE": "PO-12345",
    "TEST_DATE": "2026-03-30",
    "VALVE_SIZE": '2"',
    "MANUFACTURER": "Leser",
    "MODEL": "Type 441",
    "DISCHARGE_TYPE": "Atmosphere",
    "SET_PRESSURE": "1000",
    "VALVE_SERIAL_NO": "SN-987654321",
    "BLOWDOWN_TYPE": "Non-Adjustable",
    "TEST_GAUGE_ID": "N/A",
    "GAUGE_MODEL": "Druck DPI104",
    "GAUGE_SERIAL_NO": "6113016",
    "CERTIFICATION_NO": "N/A",
    "RECALIBRATION_DUE": "2026-04-06",
    "TEST_MEDIUM": "Nitrogen",
    "AVG_LIFT_PRESSURE": "995",
    "AVG_RESEAT_PRESSURE": "920",
    "FLOW_RATE": "450",
    "BLOWDOWN_PERCENT": "7.5",
    "ALLOWABLE_BLOWDOWN": "<15",
    "LEAK_TEST_RESULT": "PASS",
    "SET_PRESSURE_TEST_RESULT": "PASS",
    "BLOWDOWN_TEST_RESULT": "PASS",
    "OVERALL_RESULT": "PASS",
    "COMMENTS": "Valve has been tested and found to be within the +/- 3% required tolerance of the marked set pressure.",
    "TECHNICIAN_NAME": "Glenn Scatchard",
    "SIGNATURE_DATE": "2026-03-30",
}


PRESSURE_VESSEL_SAMPLE_DATA = {
    "CLIENT_COMPANY": "XYZ Industries Ltd",
    "CLIENT_ADDRESS": "456 Factory Lane, Perth WA 6000",
    "REPORT_DATE": "2026-03-30",
    "JOB_NUMBER": "WAC-2024-002",
    "REFERENCE": "PV-INSPECT-001",
    "MANUFACTURER": "Air Receiver Co",
    "MODEL": "AR-500",
    "SERIAL_NO": "SN-123456789",
    "TEST_NUMBER": "TV-002",
    "CONTACT": "Jane Doe",
    "PO_REFERENCE": "PO-67890",
    "ASSET_NO": "ASSET-001",
    "TEST_DATE": "2026-03-30",
    "UNIT_TYPE": "Air Receiver",
    "MANUFACTURED_DATE": "2020-01-15",
    "LOCATION": "Itinerant",
    "DESIGN_CODE": "AS 1210-3",
    "CORROSION_ALLOWANCE": "1.0",
    "VESSEL_TYPE": "Vertical",
    "DESIGN_REG_NO": "DR-12345",
    "WORKSAFE_REG_NO": "WR-98765",
    "COMMISSIONING_DATE": "2020-03-01",
    "DESIGN_PRESSURE": "1200",
    "OPERATING_PRESSURE": "1000",
    "TEST_PRESSURE": "1320",
    "HAZARD_LEVEL": "Level C",
    "TANK_CAPACITY": "500L",
    "HYDROSTATIC_TEST_DATE": "2020-02-01",
    "DESIGN_TEMP": "150",
    "AMBIENT_TEMP": "25",
    "SHELL_LENGTH": "1200",
    "SHELL_DIAMETER": "600",
    "NEXT_INSPECTION_DUE": "2027-03-30",
    "OVERALL_RESULT": "PASS",
    "COMMENTS": "No obvious visual defects at the time of inspection, the vessel is suitable for continued use in this application.",
    "TECHNICIAN_NAME": "Glenn Scatchard",
    "SIGNATURE_DATE": "2026-03-30",
}


def main() -> None:
    print("=" * 70)
    print("Word Template PDF Generator")
    print("=" * 70)
    print()

    ensure_safety_valve_template(refresh=True)
    ensure_pressure_vessel_template(refresh=True)

    print("Generating Safety Valve report from client template...")
    safety_output = generate_safety_valve_pdf(SAFETY_VALVE_SAMPLE_DATA, "test_output/Safety_Valve_Report_Filled.pdf")
    print(f"  Saved: {safety_output}")

    print("\nGenerating Pressure Vessel report from client template...")
    pressure_output = generate_pressure_vessel_pdf(PRESSURE_VESSEL_SAMPLE_DATA, "test_output/Pressure_Vessel_Report_Filled.pdf")
    print(f"  Saved: {pressure_output}")

    print()
    print("=" * 70)
    print("Done!")
    print("=" * 70)


if __name__ == "__main__":
    main()
