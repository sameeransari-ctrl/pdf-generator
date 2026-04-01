# from flask import Flask, render_template, request, send_file
# from docx import Document
# import os
# import uuid
# import subprocess



# app = Flask(__name__)

# TEMPLATE_PATH = "files/Safety Valve Report SV.docx"
# OUTPUT_DIR = "output"

# os.makedirs(OUTPUT_DIR, exist_ok=True)


# # ✅ 🔥 SMART REPLACE (keeps formatting + handles split placeholders)
# def replace_text(doc, key, value):

#     def process_paragraph(para):
#         if key in para.text:

#             runs = para.runs

#             # ✅ Try normal run replace (best case - keeps formatting)
#             for run in runs:
#                 if key in run.text:
#                     run.text = run.text.replace(key, value)
#                     return

#             # 🔥 Fallback for split placeholders
#             full_text = "".join(run.text for run in runs)

#             if key in full_text:
#                 new_text = full_text.replace(key, value)

#                 # Keep formatting of first run
#                 runs[0].text = new_text

#                 # Clear remaining runs
#                 for i in range(1, len(runs)):
#                     runs[i].text = ""

#     # Paragraphs
#     for para in doc.paragraphs:
#         process_paragraph(para)

#     # Tables
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for para in cell.paragraphs:
#                     process_paragraph(para)


# # ✅ LibreOffice conversion (your tested method)
# def convert_to_pdf(input_file):
#     subprocess.run([
#         r"C:\Program Files\LibreOffice\program\soffice.exe",
#         "--headless",
#         "--convert-to", "pdf",
#         "--outdir", OUTPUT_DIR,
#         input_file
#     ], check=True)


# @app.route("/", methods=["GET", "POST"])
# def index():
#     if request.method == "POST":

#         company = request.form.get("company", "")
#         address = request.form.get("address", "")
#         job_number = request.form.get("job_number", "")
#         reference = request.form.get("reference", "")
#         report_date = request.form.get("report_date", "")
#         test_number = request.form.get("test_number", "")
#         site_contact = request.form.get("site_contact", "")
#         reference_number = request.form.get("reference_number", "")
#         test_date = request.form.get("test_date", "")
#         size = request.form.get("size", "") 
#         manufacture = request.form.get("manufacture", "") 
#         modelsv = request.form.get("modelsv", "")
#         kpa = request.form.get("kpa", "")
#         vsn = request.form.get("vsn", "")
#         aflp = request.form.get("aflp", "")
#         afrp = request.form.get("afrp", "")
#         flow_r = request.form.get("flow_r", "")
#         blowdown = request.form.get("blowdown", "")
#         date = request.form.get("date", "")

#         doc = Document(TEMPLATE_PATH)

#         # 🔥 Replace all placeholders
#         replace_text(doc, "{{company}}", company)
#         replace_text(doc, "{{address}}", address)
#         replace_text(doc, "{{job_number}}", job_number)
#         replace_text(doc, "{{reference}}", reference)
#         replace_text(doc, "{{report_date}}", report_date)
#         replace_text(doc, "{{test_number}}", test_number)
#         replace_text(doc, "{{site_contact}}", site_contact)
#         replace_text(doc, "{{reference_number}}", reference_number)
#         replace_text(doc, "{{test_date}}", test_date)
#         replace_text(doc, "{{size}}", size)
#         replace_text(doc, "{{manufacture}}", manufacture)
#         replace_text(doc, "{{modelsv}}", modelsv)
#         replace_text(doc, "{{kpa}}", kpa)
#         replace_text(doc, "{{vsn}}", vsn)
#         replace_text(doc, "{{aflp}}", aflp)
#         replace_text(doc, "{{afrp}}", afrp)
#         replace_text(doc, "{{flow_r}}", flow_r)
#         replace_text(doc, "{{blowdown}}", blowdown)
#         replace_text(doc, "{{date}}", date)
        

#         # Generate filename
#         filename = str(uuid.uuid4())

#         docx_path = os.path.join(OUTPUT_DIR, f"{filename}.docx")
#         pdf_path = os.path.join(OUTPUT_DIR, f"{filename}.pdf")

#         # ✅ Save DOCX
#         doc.save(docx_path)

#         # ✅ Convert to PDF
#         convert_to_pdf(docx_path)

#         # ✅ Return PDF
#         return send_file(pdf_path, as_attachment=True)

#     return render_template("form.html")


# if __name__ == "__main__":
#     app.run(debug=True)









from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Inches
from werkzeug.utils import secure_filename
import os
import uuid
import subprocess
from datetime import datetime

app = Flask(__name__)

TEMPLATE_PATH = "files/Safety Valve Report SV.docx"
OUTPUT_DIR = "output"

os.makedirs(OUTPUT_DIR, exist_ok=True)


# ✅ 🔥 SMART REPLACE (keeps formatting + handles split placeholders)
def replace_text(doc, key, value):

    def process_paragraph(para):
        if key in para.text:

            runs = para.runs

            # ✅ Try normal run replace
            for run in runs:
                if key in run.text:
                    run.text = run.text.replace(key, value)
                    return

            # 🔥 Fallback (split placeholders)
            full_text = "".join(run.text for run in runs)

            if key in full_text:
                new_text = full_text.replace(key, value)

                runs[0].text = new_text
                for i in range(1, len(runs)):
                    runs[i].text = ""

    # Paragraphs
    for para in doc.paragraphs:
        process_paragraph(para)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para)


# ✅ LibreOffice conversion
def convert_to_pdf(input_file):
    subprocess.run([
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        "--headless",
        "--convert-to", "pdf",
        "--outdir", OUTPUT_DIR,
        input_file
    ], check=True)


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":

        # 🔹 Form Data
        company = request.form.get("company", "")
        address = request.form.get("address", "")
        job_number = request.form.get("job_number", "")
        reference = request.form.get("reference", "")
        # report_date = request.form.get("report_date", "")
        report_date = format_date(request.form.get("report_date", ""))
        test_number = request.form.get("test_number", "")
        site_contact = request.form.get("site_contact", "")
        reference_number = request.form.get("reference_number", "")
        # test_date = request.form.get("test_date", "")
        test_date = format_date(request.form.get("test_date", ""))
        size = request.form.get("size", "")
        manufacture = request.form.get("manufacture", "")
        modelsv = request.form.get("modelsv", "")
        kpa = request.form.get("kpa", "")
        vsn = request.form.get("vsn", "")
        aflp = request.form.get("aflp", "")
        afrp = request.form.get("afrp", "")
        flow_r = request.form.get("flow_r", "")
        blowdown = request.form.get("blowdown", "")
        # date = request.form.get("date", "")
        date = format_date(request.form.get("date", ""))

        doc = Document(TEMPLATE_PATH)

        # 🔥 Replace placeholders
        replace_text(doc, "{{company}}", company)
        replace_text(doc, "{{address}}", address)
        replace_text(doc, "{{job_number}}", job_number)
        replace_text(doc, "{{reference}}", reference)
        replace_text(doc, "{{report_date}}", report_date)
        replace_text(doc, "{{test_number}}", test_number)
        replace_text(doc, "{{site_contact}}", site_contact)
        replace_text(doc, "{{reference_number}}", reference_number)
        replace_text(doc, "{{test_date}}", test_date)
        replace_text(doc, "{{size}}", size)
        replace_text(doc, "{{manufacture}}", manufacture)
        replace_text(doc, "{{modelsv}}", modelsv)
        replace_text(doc, "{{kpa}}", kpa)
        replace_text(doc, "{{vsn}}", vsn)
        replace_text(doc, "{{aflp}}", aflp)
        replace_text(doc, "{{afrp}}", afrp)
        replace_text(doc, "{{flow_r}}", flow_r)
        replace_text(doc, "{{blowdown}}", blowdown)
        replace_text(doc, "{{date}}", date)

        # 🔥 Handle Image Uploads
        images = request.files.getlist("images")

        if images and any(img.filename != "" for img in images):
            doc.add_page_break()
            doc.add_heading("Attachments", level=1)

            for img in images:
                if img and img.filename != "":
                    filename_img = secure_filename(img.filename)
                    img_path = os.path.join(OUTPUT_DIR, f"{uuid.uuid4()}_{filename_img}")

                    img.save(img_path)

                    # Add image to doc
                    doc.add_picture(img_path, width=Inches(5.5))
                    doc.add_paragraph("")  # spacing

        # 🔹 File paths
        filename = str(uuid.uuid4())
        docx_path = os.path.join(OUTPUT_DIR, f"{filename}.docx")
        pdf_path = os.path.join(OUTPUT_DIR, f"{filename}.pdf")

        # ✅ Save DOCX
        doc.save(docx_path)

        # ✅ Convert to PDF
        convert_to_pdf(docx_path)

        # ✅ Return PDF
        return send_file(pdf_path, as_attachment=True)

    return render_template("form.html")



def format_date(value):
    if not value:
        return ""
    try:
        # Convert YYYY-MM-DD → DD/MM/YYYY
        return datetime.strptime(value, "%Y-%m-%d").strftime("%d/%m/%Y")
    except:
        return value  # fallback if already formatted


# if __name__ == "__main__":
#     app.run(debug=True)


if __name__ == "__main__":
    import os
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)