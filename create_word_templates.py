#!/usr/bin/env python3
"""
Create placeholder-based Word templates from the client source documents.
"""

from __future__ import annotations

import os
from typing import Dict, Iterable

from docx import Document


SAFETY_VALVE_SOURCE = "client_doc/Safety Valve Report SV (1).docx"
SAFETY_VALVE_TEMPLATE = "templates/Safety_Valve_Client_Template.docx"

PRESSURE_VESSEL_SOURCE = "client_doc/Pressure Vessel Inspection Test Report PV (1).docx"
PRESSURE_VESSEL_TEMPLATE = "templates/Pressure_Vessel_Client_Template.docx"


def _replace_paragraph_text(paragraph, replacements: Dict[str, str]) -> bool:
    text = paragraph.text
    updated = text

    for old, new in sorted(replacements.items(), key=lambda item: len(item[0]), reverse=True):
        if old in updated:
            updated = updated.replace(old, new)

    if updated == text:
        return False

    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(updated)

    return True


def _iter_all_paragraphs(doc: Document) -> Iterable:
    seen = set()

    for paragraph in doc.paragraphs:
        paragraph_id = id(paragraph)
        if paragraph_id not in seen:
            seen.add(paragraph_id)
            yield paragraph

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph_id = id(paragraph)
                    if paragraph_id not in seen:
                        seen.add(paragraph_id)
                        yield paragraph


def _create_template(source_path: str, output_path: str, replacements: Dict[str, str]) -> str:
    doc = Document(source_path)

    for paragraph in _iter_all_paragraphs(doc):
        _replace_paragraph_text(paragraph, replacements)

    doc.save(output_path)
    return output_path


def _set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _clear_paragraph(paragraph) -> None:
    _set_paragraph_text(paragraph, "")


def create_safety_valve_template() -> str:
    """Create a Safety Valve template from the client document."""
    replacements = {
        "Client / Company:": "Client / Company: {{CLIENT_COMPANY}}",
        "Address:": "Address: {{CLIENT_ADDRESS}}",
        "Report Date:": "Report Date: {{REPORT_DATE}}",
        "Job Number:": "Job Number: {{JOB_NUMBER}}",
        "Reference:": "Reference: {{REFERENCE}}",
        "Test Number:": "Test Number: {{TEST_NUMBER}}",
        "Site Contact:": "Site Contact: {{CONTACT}}",
        "PO / Reference No:": "PO / Reference No: {{PO_REFERENCE}}",
        "Test Date:": "Test Date: {{TEST_DATE}}",
        "Size:": "Size: {{VALVE_SIZE}}",
        "Manufacturer:": "Manufacturer: {{MANUFACTURER}}",
        "Leak Test : PASS Set Pressure Test: PASS Blowdown % : PASS": (
            "Leak Test : {{LEAK_TEST_RESULT}}\n"
            "Set Pressure Test: {{SET_PRESSURE_TEST_RESULT}}\n"
            "Blowdown % : {{BLOWDOWN_TEST_RESULT}}"
        ),
        "Overall Result : PASS": "Overall Result : {{OVERALL_RESULT}}",
        "Valve has been tested and found to be within the +-3% required tolerance of the marked set pressure": "{{COMMENTS}}",
        "Technician: Glenn Scatchard": "Technician: {{TECHNICIAN_NAME}}",
    }
    output_path = _create_template(SAFETY_VALVE_SOURCE, SAFETY_VALVE_TEMPLATE, replacements)

    doc = Document(output_path)
    table = doc.tables[0]
    _set_paragraph_text(
        table.cell(2, 0).paragraphs[0],
        "\nClient / Company: {{CLIENT_COMPANY}}\nAddress: {{CLIENT_ADDRESS}}\nTest Number: {{TEST_NUMBER}}",
    )
    _set_paragraph_text(
        table.cell(2, 1).paragraphs[0],
        "\nSite Contact: {{CONTACT}}\nPO / Reference No: {{PO_REFERENCE}}\nTest Date: {{TEST_DATE}}",
    )
    _set_paragraph_text(
        table.cell(3, 0).paragraphs[1],
        "Size: {{VALVE_SIZE}}\nManufacturer: {{MANUFACTURER}}\nModel: {{MODEL}}\n",
    )
    _set_paragraph_text(
        table.cell(3, 1).paragraphs[0],
        "\n\nDischarge Type: {{DISCHARGE_TYPE}}\nSet Pressure (kPa): {{SET_PRESSURE}}\n",
    )
    _set_paragraph_text(
        table.cell(3, 3).paragraphs[0],
        "\n\nValve Serial No: {{VALVE_SERIAL_NO}}\nBlowdown Type: {{BLOWDOWN_TYPE}}",
    )
    _set_paragraph_text(
        table.cell(5, 0).paragraphs[0],
        "\nTest Gauge ID: {{TEST_GAUGE_ID}}\nGauge Model: {{GAUGE_MODEL}}\nGauge Serial No: {{GAUGE_SERIAL_NO}}\n",
    )
    _set_paragraph_text(
        table.cell(5, 4).paragraphs[0],
        "\nCertification No: {{CERTIFICATION_NO}}\nRecalibration Due: {{RECALIBRATION_DUE}}\nTest Medium: {{TEST_MEDIUM}}",
    )
    _set_paragraph_text(
        table.cell(6, 0).paragraphs[1],
        "Average Final Lift Pressure (kPa): {{AVG_LIFT_PRESSURE}}\n"
        "Average Final Reseat Pressure (kPa): {{AVG_RESEAT_PRESSURE}} Flow Rate (scfm) : {{FLOW_RATE}}",
    )
    _set_paragraph_text(
        table.cell(6, 2).paragraphs[0],
        "\n\nBlowdown %: {{BLOWDOWN_PERCENT}}\nAllowable Blowdown %: {{ALLOWABLE_BLOWDOWN}}",
    )
    _set_paragraph_text(
        table.cell(6, 5).paragraphs[0],
        "\n\nLeak Test: {{LEAK_TEST_RESULT}}\nSet Pressure Test: {{SET_PRESSURE_TEST_RESULT}}\nBlowdown %: {{BLOWDOWN_TEST_RESULT}}\n",
    )
    doc.paragraphs[21].text = "Date: {{SIGNATURE_DATE}}"
    doc.save(output_path)
    return output_path


def create_pressure_vessel_template() -> str:
    """Create a Pressure Vessel template from the client document."""
    replacements = {
        "Client / Company:": "Client / Company: {{CLIENT_COMPANY}}",
        "Address:": "Address: {{CLIENT_ADDRESS}}",
        "Report Date:": "Report Date: {{REPORT_DATE}}",
        "Job Number:": "Job Number: {{JOB_NUMBER}}",
        "Reference:": "Reference: {{REFERENCE}}",
        "Test Number:": "Test Number: {{TEST_NUMBER}}",
        "Site Contact:": "Site Contact: {{CONTACT}}",
        "PO / Reference No:": "PO / Reference No: {{PO_REFERENCE}}",
        "Asset No:": "Asset No: {{ASSET_NO}}",
        "Test Date:": "Test Date: {{TEST_DATE}}",
        "Unit Type:": "Unit Type: {{UNIT_TYPE}}",
        "Next Scheduled Inspection Due:": "Next Scheduled Inspection Due: {{NEXT_INSPECTION_DUE}}",
        "Compliance with AS/NZS 3788:2024: PASS ☐ FAIL": "Compliance with AS/NZS 3788:2024: {{OVERALL_RESULT}}",
        "No obvious visual defects at the time of inspection, the vessel is suitable for continued use in this application": "{{COMMENTS}}",
        ": Glenn Scatchard": ": {{TECHNICIAN_NAME}}",
    }
    output_path = _create_template(PRESSURE_VESSEL_SOURCE, PRESSURE_VESSEL_TEMPLATE, replacements)

    doc = Document(output_path)
    table = doc.tables[0]

    _set_paragraph_text(
        table.cell(2, 0).paragraphs[0],
        "\nClient / Company: {{CLIENT_COMPANY}}\nAddress: {{CLIENT_ADDRESS}}\n"
        "Manufacturer: {{MANUFACTURER}}\nModel: {{MODEL}}\nSerial No: {{SERIAL_NO}}\nTest Number: {{TEST_NUMBER}}",
    )
    _set_paragraph_text(
        table.cell(2, 1).paragraphs[0],
        "\nSite Contact: {{CONTACT}}\nPO / Reference No: {{PO_REFERENCE}}\nAsset No: {{ASSET_NO}}\n"
        "Test Date: {{TEST_DATE}}\nUnit Type: {{UNIT_TYPE}}",
    )
    _set_paragraph_text(
        table.cell(3, 0).paragraphs[1],
        "Manufacturer: {{VESSEL_MANUFACTURER}}\nManufactured Date: {{MANUFACTURED_DATE}}\nLocation: {{LOCATION}}\n",
    )
    _set_paragraph_text(
        table.cell(3, 1).paragraphs[1],
        "Serial No: {{VESSEL_SERIAL_NO}}\nDesign Code: {{DESIGN_CODE}}\nCorrosion Allowance (mm): {{CORROSION_ALLOWANCE}}",
    )
    _set_paragraph_text(
        table.cell(4, 0).paragraphs[1],
        "Type: {{VESSEL_TYPE}}\nDesign Registration No: {{DESIGN_REG_NO}}\nWorkSafe Registration No: {{WORKSAFE_REG_NO}}\n"
        "Commissioning Date: {{COMMISSIONING_DATE}}\nDesign Pressure (Kpa): {{DESIGN_PRESSURE}}\n"
        "Operating Pressure (Kpa): {{OPERATING_PRESSURE}}\nTest Pressure (Kpa): {{TEST_PRESSURE}}",
    )
    _set_paragraph_text(table.cell(4, 0).paragraphs[2], "Hazard Level (as per AS/NZS 3788:2024): {{HAZARD_LEVEL}}")
    _set_paragraph_text(
        table.cell(4, 1).paragraphs[1],
        "Tank Capacity: {{TANK_CAPACITY}}\nHydrostatic Test Date: {{HYDROSTATIC_TEST_DATE}}\n"
        "Design Temperature (Degrees): {{DESIGN_TEMP}}\nAmbient Air Temperature (Degrees): {{AMBIENT_TEMP}}\n"
        "Shell Length (mm): {{SHELL_LENGTH}}\nShell Diameter (mm): {{SHELL_DIAMETER}}",
    )
    _set_paragraph_text(
        table.cell(5, 0).paragraphs[5],
        "Inspection Type:\n\n{{INSPECTION_TYPE}}",
    )
    _set_paragraph_text(
        table.cell(6, 0).paragraphs[2],
        "LHS/Lid - Spec: {{LHS_LID_SPEC}}    Actual: {{LHS_LID_ACTUAL}}\n"
        "RHS/Base - Spec: {{RHS_BASE_SPEC}}    Actual: {{RHS_BASE_ACTUAL}}\n"
        "Shell - Spec: {{SHELL_SPEC}}    Actual: {{SHELL_ACTUAL}}\n",
    )
    _set_paragraph_text(
        table.cell(6, 0).paragraphs[4],
        "External Surface Condition:\n{{EXTERNAL_SURFACE_CONDITION}}\nComments: {{EXTERNAL_SURFACE_COMMENTS}}",
    )
    _set_paragraph_text(
        table.cell(6, 0).paragraphs[5],
        "Internal Surface Condition (if applicable):\n{{INTERNAL_SURFACE_CONDITION}}\nComments: {{INTERNAL_SURFACE_COMMENTS}}",
    )
    _set_paragraph_text(
        table.cell(6, 0).paragraphs[6],
        "Corrosion / Pitting Observed:\n{{CORROSION_OBSERVED}}\nLocation: {{CORROSION_LOCATION}}",
    )
    _set_paragraph_text(
        table.cell(6, 0).paragraphs[7],
        "Dents / Deformation:\n{{DENTS_DEFORMATION}}\nDetails: {{DENTS_DETAILS}}",
    )
    _set_paragraph_text(
        table.cell(6, 0).paragraphs[8],
        "Pressure Gauge Working:\n{{PRESSURE_GAUGE_WORKING}}",
    )
    _clear_paragraph(table.cell(6, 0).paragraphs[9])
    _set_paragraph_text(
        table.cell(6, 1).paragraphs[0],
        "Paint / Coating Condition:\n{{PAINT_COATING_CONDITION}}",
    )
    _set_paragraph_text(
        table.cell(6, 1).paragraphs[1],
        "Base / Mounts / Supports:\n{{BASE_MOUNTS_CONDITION}}\nComments: {{BASE_MOUNTS_COMMENTS}}",
    )
    _set_paragraph_text(
        table.cell(6, 1).paragraphs[2],
        "Nameplate:\n{{NAMEPLATE_STATUS}}",
    )
    _set_paragraph_text(
        table.cell(6, 1).paragraphs[3],
        "Drain Valve Operation:\n{{DRAIN_VALVE_OPERATION}}",
    )
    _set_paragraph_text(
        table.cell(6, 1).paragraphs[4],
        "Safety Valve Fitted:\n{{SAFETY_VALVE_FITTED}}\nSee Report No: {{SAFETY_VALVE_REPORT_NO}}",
    )
    _set_paragraph_text(
        table.cell(6, 1).paragraphs[5],
        "Safety Valve Details:\n{{SAFETY_VALVE_DETAILS}}",
    )
    for paragraph in table.cell(6, 1).paragraphs[7:]:
        _clear_paragraph(paragraph)
    for paragraph in table.cell(6, 2).paragraphs[7:]:
        _clear_paragraph(paragraph)
    _set_paragraph_text(
        table.cell(8, 0).paragraphs[1],
        "Overall Vessel Condition:\n{{OVERALL_VESSEL_CONDITION}}",
    )
    _set_paragraph_text(
        table.cell(8, 0).paragraphs[2],
        "Compliance with AS/NZS 3788:2024:\n{{OVERALL_RESULT}}",
    )
    _set_paragraph_text(
        table.cell(8, 0).paragraphs[3],
        "Documentation:\n\nDesign Reports Sighted: {{DESIGN_REPORTS_SIGHTED}}\n"
        "Next Scheduled Inspection Due: {{NEXT_INSPECTION_DUE}}\nAssociated Documentation: {{ASSOCIATED_DOCUMENTATION}}",
    )
    _set_paragraph_text(doc.tables[1].cell(2, 0).paragraphs[0], "Date: {{SIGNATURE_DATE}}")
    doc.save(output_path)
    return output_path


def create_all_templates() -> None:
    """Create both client-based templates."""
    os.makedirs("templates", exist_ok=True)
    create_safety_valve_template()
    create_pressure_vessel_template()


def main() -> None:
    print("=" * 70)
    print("Creating Word Document Templates with Placeholders")
    print("=" * 70)
    print()

    create_all_templates()

    print("Templates saved in 'templates/' folder:")
    print("  - Safety_Valve_Client_Template.docx")
    print("  - Pressure_Vessel_Client_Template.docx")
    print()
    print("Next step: Use fill_templates.py or desktop_pdf_app.py to generate PDFs")


if __name__ == "__main__":
    main()
